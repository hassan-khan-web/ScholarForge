import os
from serpapi import GoogleSearch
from docx import Document
from docx.shared import Inches, Pt, RGBColor
import httpx
from bs4 import BeautifulSoup
import json
import re
import matplotlib
matplotlib.use('Agg') 
import matplotlib.pyplot as plt
import pandas as pd
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
import fitz 

from report_formats import get_template_instructions

SMART_MODEL = "qwen/qwen3-coder:free"
BACKUP_MODEL = "mistralai/devstral-2512:free"

SEARCH_RESULTS_COUNT = 10
MAX_RESULTS_TO_SCRAPE = 4
WORDS_PER_PAGE = 450

def clean_ai_output(text: str) -> str:
    if not text: return ""
    text = re.sub(r'<think>.*?</think>', '', text, flags=re.DOTALL)
    text = re.sub(r'^```\w*\s*', '', text, flags=re.MULTILINE)
    text = re.sub(r'^```\s*$', '', text, flags=re.MULTILINE)
    return text.strip()

def clean_section_output(text: str, section_title: str) -> str:
    if not text: return ""
    text = clean_ai_output(text)
    lines = text.split('\n')
    while lines and not lines[0].strip(): lines.pop(0)
    if not lines: return ""
    
    first_line = lines[0].strip().lower()
    clean_title = section_title.lower().replace('#', '').strip()
    clean_first = first_line.replace('#', '').strip()
    
    if clean_title in clean_first or clean_first in clean_title:
        return "\n".join(lines[1:]).strip()
    return text.strip()

def call_llm(target_model: str, system_prompt: str, user_prompt: str, temp: float = 0.4, attempt: int = 1) -> str:
    current_model = target_model
    if attempt == 2:
        current_model = BACKUP_MODEL
        print(f"    >>> Model Switch: {current_model}")
    elif attempt > 2:
        return "Error: AI models unavailable."

    try:
        api_key = os.environ.get("OPENROUTER_API_KEY")
        timeout = 120.0
        
        system_prompt += " Output raw Markdown only. No code blocks."

        with httpx.Client(timeout=timeout) as client:
            response = client.post(
                url="https://openrouter.ai/api/v1/chat/completions",
                headers={
                    "Authorization": f"Bearer {api_key}", 
                    "Content-Type": "application/json",
                    "HTTP-Referer": "http://localhost:5000",
                    "X-Title": "ScholarForge"
                },
                json={
                    "model": current_model,
                    "messages": [
                        {"role": "system", "content": system_prompt}, 
                        {"role": "user", "content": user_prompt}
                    ],
                    "temperature": temp,
                    "max_tokens": 5000 
                }
            )
            if response.status_code != 200:
                print(f"    [!] AI Error ({current_model}): {response.status_code}")
                return call_llm(target_model, system_prompt, user_prompt, temp, attempt + 1)
                
            return clean_ai_output(response.json()['choices'][0]['message']['content'])
    except Exception as e:
        print(f"    [!] Exception ({current_model}): {e}")
        return call_llm(target_model, system_prompt, user_prompt, temp, attempt + 1)


def extract_text_from_multiple_pdfs(file_bytes_list: list) -> str:
    """Feature: Extract text from MULTIPLE uploaded PDFs"""
    combined_text = "\n\n--- USER UPLOADED DOCUMENTS ---\n"
    
    try:
        for idx, file_bytes in enumerate(file_bytes_list):
            with fitz.open(stream=file_bytes, filetype="pdf") as doc:
                doc_text = ""
                for i, page in enumerate(doc):
                    if i > 25: break
                    doc_text += page.get_text()
                
                combined_text += f"\n[Document {idx+1} Content]:\n{doc_text[:15000]}\n" 
        
        combined_text += "\n------------------------------\n"
        return combined_text
    except Exception as e:
        print(f"PDF Error: {e}")
        return ""

def _get_article_text(url: str) -> str:
    try:
        headers = {'User-Agent': 'Mozilla/5.0'}
        with httpx.Client(timeout=10.0, follow_redirects=True) as client:
            response = client.get(url, headers=headers)
        if response.status_code != 200: return ""
        
        if "application/pdf" in response.headers.get("Content-Type", "") or url.endswith(".pdf"):
            return "" 
            
        soup = BeautifulSoup(response.text, 'lxml')
        for tag in soup(['script', 'style', 'nav', 'footer', 'aside']): tag.decompose()
        return soup.get_text(separator='\n', strip=True)[:5000] # Increased scrape limit
    except: return ""

def get_search_results(query: str, max_results: int = SEARCH_RESULTS_COUNT) -> str:
    """Feature: Structured Source Verification"""
    try:
        api_key = os.environ.get("SERPAPI_KEY") 
        if not api_key: return "Error: SERPAPI_KEY not set."
        params = {
            "q": query,
            "location": "US",
            "hl": "en",
            "gl": "us",
            "num": 5,
            "engine": "google",
            "api_key": api_key
        }
        search = GoogleSearch(params)
        results = search.get_dict()
        
        formatted_output = "--- VERIFIED SOURCES ---\n"
        
        if "organic_results" in results:
            for i, result in enumerate(results["organic_results"]):
                if i >= MAX_RESULTS_TO_SCRAPE: break
                
                url = result.get("link", "")
                title = result.get('title', 'Unknown Title')
                snippet = result.get("snippet", "")
                
                full_content = ""
                if url:
                    raw = _get_article_text(url)
                    if raw: full_content = f"\nEXTRACT: {raw}"
                
                formatted_output += f"SOURCE [{i+1}]\nTitle: {title}\nURL: {url}\nSummary: {snippet}{full_content}\n\n"
                
        return formatted_output
    except Exception as e: return f"Search Error: {e}"


def recursive_gap_analysis(section_title: str, existing_summary: str, topic: str) -> str:
    """Feature: Recursive Research. Checks if we need more info."""
    print(f"    > Analyzing gap for: {section_title}")
    prompt = (
        f"We are writing a report on '{topic}'.\n"
        f"Current Section: '{section_title}'\n"
        f"Available Data Summary: {existing_summary[:3000]}\n\n"
        "DECISION: Do we have specific enough data to write a detailed 600-word section with stats and tables on this specific sub-topic?\n"
        "If YES, output 'PASS'.\n"
        "If NO, output a Google Search Query to find the missing specific info."
    )
    decision = call_llm(SMART_MODEL, "You are a Research Director.", prompt, temp=0.1)
    
    if "PASS" in decision or len(decision) > 100:
        return "" 
    
    new_query = decision.strip().replace('"', '')
    print(f"    >>> RECURSIVE SEARCH TRIGGERED: {new_query}")
    return get_search_results(new_query, max_results=2)

def assess_search_need(query: str, existing_context: str) -> str:
    """Feature: Check if we actually need to search the web."""
    print(f"    > Assessing search need for: {query}")
    prompt = (
        f"Query: '{query}'\n"
        f"Existing Context Length: {len(existing_context)} chars\n"
        f"Existing Context Preview: {existing_context[:1000]}\n\n"
        "DECISION: To write a high-quality, detailed report on this, do we STRICTLY need external live web search data?\n"
        "Criteria:\n"
        "- If it is a well-known topic (history, science, standard concepts) or purely creative -> NO.\n"
        "- If the provided Context answers it -> NO.\n"
        "- If it requires REAL-TIME news, specific recent data (post-2023), or obscure info -> YES.\n\n"
        "OUTPUT:\n"
        "- If NO (we can skip search): Output 'SKIP_SEARCH'.\n"
        "- If YES (we need search): Output a specific, optimized Google Search Query."
    )
    decision = call_llm(SMART_MODEL, "You are a Research Director.", prompt, temp=0.1)
    
    clean_decision = decision.strip().replace('"', '')
    if 'SKIP_SEARCH' in clean_decision:
        return 'SKIP_SEARCH'
    return clean_decision

def generate_summary(search_content: str, topic: str, user_pdf_text: str = "") -> str:
    context = search_content
    if user_pdf_text:
        context = user_pdf_text + "\n\n" + search_content
        
    return call_llm(
        SMART_MODEL,
        "You are a Senior Research Analyst.",
        f"Topic: {topic}\n\nData:\n{context[:35000]}\n\nTask: Synthesize a master summary of key facts, numbers, and sources. Group them by themes.\nIMPORTANT: If the Data seems empty or insufficient, rely on your extensive INTERNAL KNOWLEDGE to generate the summary."
    )

def generate_outline(topic: str, summary: str, format_type: str, target_pages: int) -> list:
    format_data = get_template_instructions(format_type, target_pages)
    
    target_count = format_data['target_sections']
    
    prompt = (
        f"Topic: {topic}\n"
        f"Tier Target: {target_count} sections exactly.\n"
        f"Logic: {format_data['template_text']}\n"
        f"Context: {summary[:3000]}\n\n"
        "TASK: Generate the JSON outline.\n"
        "RULES:\n"
        "1. Titles MUST be engaging (e.g. 'The Quantum Leap' instead of 'Introduction').\n"
        "2. Return exactly the number of sections requested.\n"
        "Output: A JSON list of strings ONLY. Example: [\"1. The Awakening\", \"2. Market Forces\"]"
    )
    content = call_llm(SMART_MODEL, "Return JSON only.", prompt, temp=0.3)
    match = re.search(r'\[.*\]', content.replace('\n', ' '), re.DOTALL)
    
    if match: 
        outline = json.loads(match.group(0))
        return outline[:target_count] if len(outline) > target_count else outline
        
    return ["1. Executive Overview", "2. Core Analysis", "3. Strategic Implications", "4. Conclusion"]

def write_section(section_title: str, topic: str, summary: str, full_report_context: str, word_limit: int) -> str:
    new_data = recursive_gap_analysis(section_title, summary, topic)
    
    combined_data = summary
    if new_data:
        combined_data = new_data + "\n\n" + summary 
        
    base_prompt = (
        f"Write the section '{section_title}' for the report '{topic}'.\n"
        f"Data Source:\n{combined_data[:20000]}\n\n"
        f"Length Target: {word_limit} words.\n\n"
        "FORMATTING RULES (STRICT):\n"
        "1. HEADER: Use the section title as # H1.\n"
        "2. SUB-HEADERS: Use ### H3 for sub-themes. Do NOT use generic names.\n"
        "3. TABLES: You MUST include at least one Markdown table comparing data, pros/cons, or timelines.\n"
        "4. CITATIONS: Use [1], [2] notation corresponding to sources.\n"
        "5. TONE: Professional, dense, and analytical. Avoid fluff.\n"
        "6. CONTENT: If this is 'Standard' or higher, include a 'Real World Application' subsection."
    )
    
    content = call_llm(SMART_MODEL, "You are a Report Writer. Use Markdown Tables and Charts.", base_prompt, temp=0.4)
    return clean_section_output(content, section_title)

def generate_chart_from_data(summary: str, topic: str) -> str:
    try:
        chart_dir = "/app/static/charts"
        if not os.path.exists(chart_dir): os.makedirs(chart_dir, exist_ok=True)
        clean_name = re.sub(r'\W+', '', topic)[:15] 
        filename = f"chart_{clean_name}_{os.urandom(4).hex()}.png"
        filepath = os.path.join(chart_dir, filename)

        prompt = (
            f"Topic: {topic}\nContext: {summary[:3000]}\n"
            "Extract key numeric trends. Return JSON: {\"title\": \"...\", \"x_label\": \"...\", \"y_label\": \"...\", \"data\": [{\"label\": \"A\", \"value\": 10}]}"
        )
        content = call_llm(SMART_MODEL, "Return JSON only.", prompt, temp=0.1)
        match = re.search(r'\{.*\}', content.replace('\n', ' '), re.DOTALL)
        if not match: return None
        chart_data = json.loads(match.group(0))
        if not chart_data or 'data' not in chart_data: return None

        df = pd.DataFrame(chart_data['data'])
        fig, ax = plt.subplots(figsize=(10, 6))
        plt.style.use('ggplot')
        ax.bar(df['label'], df['value'], color='#4f46e5', alpha=0.8)
        ax.set_title(chart_data.get('title', 'Analysis'), fontsize=14, pad=20)
        ax.set_xlabel(chart_data.get('x_label', ''), fontsize=12)
        ax.set_ylabel(chart_data.get('y_label', ''), fontsize=12)
        plt.setp(ax.get_xticklabels(), rotation=45, ha='right')
        fig.tight_layout()
        fig.savefig(filepath, dpi=100)
        plt.close(fig)
        return filepath
    except: return None

def run_ai_engine_with_return(query: str, user_format: str, page_count: int = 15, pdf_bytes_list: list = None, task=None) -> tuple[str, str, str]: 
    def _update_status(message: str):
        print(message) 
        if task: task.update_state(state='PROGRESS', meta={'message': message})

    _update_status("Step 1/7: Processing Inputs...")
    
    user_pdf_text = ""
    if pdf_bytes_list:
        user_pdf_text = extract_text_from_multiple_pdfs(pdf_bytes_list)
        _update_status(f"    > Analyzed {len(pdf_bytes_list)} uploaded documents.")

    _update_status("Step 2/7: Checking Information Needs...")
    
    search_decision = assess_search_need(query, user_pdf_text)
    
    search_content = ""
    if search_decision == 'SKIP_SEARCH':
        _update_status("    > Sufficient internal/provided info. Skipping Web Search.")
        search_content = "[Internal Knowledge & User Documents Mode Active - Web Search Skipped]"
    else:
        _update_status(f"    > Web Search Required: {search_decision}")
        search_content = get_search_results(search_decision)
    
    _update_status("Step 3/7: Synthesizing Data...")
    summary = generate_summary(search_content, query, user_pdf_text)
    
    _update_status("Step 4/7: Generating Visuals...")
    chart_path = generate_chart_from_data(summary, query)
    
    _update_status("Step 5/7: Planning Structure...")
    outline = generate_outline(query, summary, user_format, page_count)

    total_words = page_count * WORDS_PER_PAGE 
    words_per_section = max(400, int(total_words / max(1, len(outline))))
    
    full_report = f"# {query.upper()}\n\n"
    for i, section in enumerate(outline):
        _update_status(f"Step 6/7: Writing Section {i+1}/{len(outline)}: {section}...")
        section_content = write_section(section, query, summary, full_report, words_per_section)
        full_report += f"\n\n## {section}\n{section_content}\n"
    
    _update_status("Step 7/7: Finalizing...")
    full_report = clean_ai_output(full_report)
    
    return search_content + "\n" + user_pdf_text, full_report, chart_path

def convert_to_txt(content, path):
    with open(path, "w", encoding="utf-8") as f: f.write(content)
    return "Success"
def convert_to_md(content, path):
    with open(path, "w", encoding="utf-8") as f: f.write(content)
    return "Success"
def convert_to_json(content, topic, path):
    data = {"topic": topic, "content": content, "generated_by": "ScholarForge"}
    with open(path, "w", encoding="utf-8") as f: json.dump(data, f, indent=4)
    return "Success"
def add_formatted_text(paragraph, text):
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
        else: paragraph.add_run(part)
def _add_markdown_table_to_docx(doc, table_block):
    lines = [l.strip() for l in table_block.strip().split('\n') if l.strip()]
    if len(lines) < 3: return
    rows = []
    for line in lines:
        if '---' in line: continue 
        cells = [c.strip().replace('**','') for c in line.strip('|').split('|')]
        rows.append(cells)
    if not rows: return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Table Grid'
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_data in enumerate(row_data):
            if c_idx < len(table.columns):
                cell = table.cell(r_idx, c_idx)
                cell.text = cell_data
                if r_idx == 0: cell.paragraphs[0].runs[0].bold = True
def convert_to_docx(content, topic, path, chart_path=None):
    doc = Document()
    doc.add_heading(topic, 0)
    if chart_path and os.path.exists(chart_path):
        try: doc.add_picture(chart_path, width=Inches(6)); doc.add_paragraph("Figure 1: Analysis", style='Caption')
        except: pass
    lines = content.split('\n')
    table_buffer = []
    in_table = False
    for line in lines:
        stripped = line.strip()
        if stripped.startswith('|') and '|' in stripped[1:]:
            in_table = True; table_buffer.append(stripped); continue
        elif in_table:
            _add_markdown_table_to_docx(doc, "\n".join(table_buffer)); table_buffer = []; in_table = False
        if not stripped: continue
        if stripped.startswith('## '): doc.add_heading(stripped.replace('## ', ''), level=2)
        elif stripped.startswith('# '): doc.add_heading(stripped.replace('# ', ''), level=1)
        elif stripped.startswith('### '): doc.add_heading(stripped.replace('### ', ''), level=3)
        elif stripped.startswith('* ') or stripped.startswith('- '): p = doc.add_paragraph(style='List Bullet'); add_formatted_text(p, stripped[2:])
        else: p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6); add_formatted_text(p, stripped)
    if in_table and table_buffer: _add_markdown_table_to_docx(doc, "\n".join(table_buffer))
    doc.save(path); return "Success"

def format_pdf_text(text):
    text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    return re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)

def convert_to_pdf(content, topic, path, chart_path=None):
    doc = SimpleDocTemplate(path, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
    styles = getSampleStyleSheet()
    story = [Paragraph(topic, styles['Title']), Spacer(1, 12)]
    if chart_path and os.path.exists(chart_path):
        try: story.append(RLImage(chart_path, width=450, height=250)); story.append(Spacer(1, 12))
        except: pass
    lines = content.split('\n')
    table_buffer = []
    in_table = False
    for line in lines:
        stripped = line.strip()
        if stripped.startswith('|') and '|' in stripped[1:]:
            in_table = True; table_buffer.append(stripped); continue
        elif in_table:
            data = [[c.strip().replace('**','') for c in row.strip('|').split('|') if '---' not in row] for row in table_buffer if '---' not in row]
            if data: story.append(Table(data, style=[('GRID', (0,0), (-1,-1), 1, colors.black)])); story.append(Spacer(1, 12))
            table_buffer = []; in_table = False
        if not stripped: continue
        clean_text = format_pdf_text(stripped)
        if stripped.startswith('##'): story.append(Paragraph(clean_text.replace('#', ''), styles['Heading2']))
        elif stripped.startswith('*') or stripped.startswith('-'): story.append(Paragraph(f"• {clean_text[2:]}", styles['Normal']))
        else: story.append(Paragraph(clean_text, styles['Normal']))
    doc.build(story); return "Success"